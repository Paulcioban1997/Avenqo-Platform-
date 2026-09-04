"""Tenant-scoped cleaning lineage, preview, and export over existing ingestion artifacts."""

from __future__ import annotations

import csv
from dataclasses import asdict
from datetime import datetime
from io import BytesIO, StringIO
import json
from pathlib import Path
from typing import Any
from uuid import UUID

from docx import Document
from openpyxl import Workbook
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from reportlab.lib import colors

from backend.app.models import DatasetStatus
from backend.app.services.company_dataset_ingestion_service import CompanyDatasetIngestionService
from shared.ai_engine.contracts import TenantContext
from shared.ai_engine.dataset_ingestion.exceptions import (
    DatasetArtifactMissingError,
    DatasetIngestionError,
)
from shared.ai_engine.dataset_ingestion.cleaning import CleaningReport
from shared.ai_engine.dataset_ingestion.quality import assess_quality


class DatasetNotReadyForExport(ValueError):
    pass


class DatasetSourceUnavailable(DatasetNotReadyForExport):
    """Le fichier original de ce dataset n'est plus lisible sur le stockage.

    Distinct de `DatasetNotReadyForExport` (qui signifie "pas encore prêt")
    afin que l'appelant (routeur/API) puisse renvoyer une raison structurée
    et actionnable ("reason": "source_artifact_missing") plutôt qu'un message
    générique masquant la vraie cause.
    """

    reason = "source_artifact_missing"


class UnsupportedDatasetExport(ValueError):
    pass


class DatasetCleaningService:
    _MAX_PREVIEW_ROWS = 100
    _REPORT_PREVIEW_ROWS = 20

    def __init__(self, ingestion: CompanyDatasetIngestionService) -> None:
        self._ingestion = ingestion

    def detail(
        self,
        tenant: TenantContext,
        dataset_id: UUID,
        *,
        offset: int = 0,
        limit: int = 25,
    ) -> dict[str, Any]:
        dataset = self._ingestion.get(tenant, dataset_id)
        version = next((item for item in dataset.versions if item.is_current), None)
        if version is None:
            raise DatasetNotReadyForExport("No current dataset version is available.")

        original_rows = self._ingestion._reload_current_version_rows(dataset)
        if not original_rows and self._ingestion._expected_row_count(dataset) > 0:
            # Le dataset avait des lignes lors de son import mais le fichier
            # original est introuvable sur le stockage actuel : ne jamais
            # afficher un faux "0 -> 0" ni une erreur générique.
            raise DatasetSourceUnavailable(
                "Le fichier original de ce dataset n'est plus disponible sur le "
                "stockage. Réimportez le fichier pour restaurer les données "
                "nettoyées."
            )
        cleaned_rows: list[dict[str, Any]] = []
        cleaning_report = self._empty_report(len(original_rows))
        mappings: dict[str, str] = {}
        quality_status: str | None = None
        quality_reasons: list[str] = []

        metadata = self._metadata(tenant.company_id, dataset.id, version.version_number)
        try:
            cleaned_rows = [
                dict(row) for row in self._ingestion.get_cleaned_rows(tenant, dataset_id)
            ]
        except DatasetIngestionError:
            try:
                self._ingestion.ensure_cleaning_artifacts(tenant, dataset)
            except DatasetArtifactMissingError as exc:
                raise DatasetSourceUnavailable(str(exc)) from exc
            metadata = self._metadata(tenant.company_id, dataset.id, version.version_number)
            try:
                cleaned_rows = [
                    dict(row) for row in self._ingestion.get_cleaned_rows(tenant, dataset_id)
                ]
            except DatasetIngestionError as exc:
                raise DatasetNotReadyForExport(str(exc)) from exc
        cleaning_report = self._report(metadata, len(original_rows), len(cleaned_rows))
        quality = assess_quality(cleaning_report)
        quality_status = quality.status.value
        quality_reasons = list(quality.reasons)
        if dataset.mapping is not None:
            mappings = dict(dataset.mapping.mapping_json.get("accepted") or {})

        bounded_limit = max(1, min(limit, self._MAX_PREVIEW_ROWS))
        summary = self._summary(
            dataset.columns_count,
            version.version_number,
            cleaning_report,
            mappings,
            metadata,
        )
        return {
            "dataset_id": str(dataset.id),
            "name": dataset.name,
            "status": self._business_status(dataset.status),
            "cleaning_status": quality_status or "configuration_required",
            "quality_reasons": quality_reasons,
            "version": version.version_number,
            "timestamp": version.created_at,
            "summary": summary,
            "original_preview": original_rows[: self._REPORT_PREVIEW_ROWS],
            "cleaned_preview": cleaned_rows[offset : offset + bounded_limit],
            "preview_offset": offset,
            "preview_limit": bounded_limit,
            "preview_total": len(cleaned_rows),
            "transformation_history": [
                {
                    "version": version.version_number,
                    "timestamp": version.created_at,
                    "summary": summary,
                }
            ],
        }

    def export(self, tenant: TenantContext, dataset_id: UUID, export_format: str) -> tuple[bytes, str, str]:
        detail = self.detail(tenant, dataset_id, limit=self._MAX_PREVIEW_ROWS)
        dataset = self._ingestion.get(tenant, dataset_id)
        rows = [dict(row) for row in self._ingestion.get_cleaned_rows(tenant, dataset_id)]
        safe_stem = Path(dataset.name).stem or "dataset"
        normalized = export_format.casefold()
        if normalized == "csv":
            return self._csv(rows), "text/csv; charset=utf-8", f"{safe_stem}-cleaned.csv"
        if normalized == "xlsx":
            return self._xlsx(rows), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", f"{safe_stem}-cleaned.xlsx"
        if normalized == "pdf":
            return self._pdf(detail, rows[: self._REPORT_PREVIEW_ROWS]), "application/pdf", f"{safe_stem}-cleaning-report.pdf"
        if normalized == "docx":
            return self._docx(detail, rows[: self._REPORT_PREVIEW_ROWS]), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", f"{safe_stem}-cleaning-report.docx"
        raise UnsupportedDatasetExport(f"Unsupported export format: {export_format}")

    @staticmethod
    def _business_status(status: DatasetStatus) -> str:
        if status in {DatasetStatus.READY, DatasetStatus.VALIDATED}:
            return "ready"
        if status == DatasetStatus.MAPPING_REQUIRED:
            return "attention_required"
        return "error"

    @staticmethod
    def _empty_report(rows: int) -> CleaningReport:
        return CleaningReport(rows, rows, 0, 0, 0, 0, 0)

    @staticmethod
    def _report(
        metadata: dict[str, Any],
        original_rows: int,
        cleaned_rows: int,
    ) -> CleaningReport:
        persisted = dict(metadata.get("cleaning_report") or {})
        return CleaningReport(
            rows_before=int(persisted.get("original_row_count", original_rows)),
            rows_after=int(persisted.get("cleaned_row_count", cleaned_rows)),
            duplicates_removed=int(persisted.get("duplicate_rows_removed", 0)),
            numeric_conversions=int(persisted.get("numeric_normalizations", 0)),
            date_conversions=int(persisted.get("date_normalizations", 0)),
            null_cells_detected=int(persisted.get("missing_values_detected", 0)),
            invalid_rows=int(persisted.get("invalid_rows_rejected", 0)),
        )

    def _metadata(self, company_id: UUID, dataset_id: UUID, version_number: int) -> dict[str, Any]:
        path = self._ingestion._storage.metadata_path(company_id, dataset_id, version_number)
        if not path.is_file():
            return {}
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            return {}

    @staticmethod
    def _summary(
        column_count: int,
        version: int,
        report: CleaningReport,
        mappings: dict[str, str],
        metadata: dict[str, Any],
    ) -> dict[str, Any]:
        persisted = dict(metadata.get("cleaning_report") or {})
        return {
            "original_row_count": report.rows_before,
            "cleaned_row_count": report.rows_after,
            "column_count": column_count,
            "columns_renamed": persisted.get("columns_renamed", {}),
            "inferred_data_types": metadata.get("inferred_data_types", {}),
            "missing_values_detected": report.null_cells_detected,
            "missing_values_corrected": 0,
            "duplicate_rows_detected": report.duplicates_removed,
            "duplicate_rows_removed": report.duplicates_removed,
            "invalid_rows_rejected": 0,
            "normalization_performed": report.numeric_conversions + report.date_conversions > 0,
            "date_normalizations": report.date_conversions,
            "numeric_normalizations": report.numeric_conversions,
            "outlier_handling": None,
            "mappings_applied": mappings,
            "dataset_version": version,
        }

    @staticmethod
    def _fieldnames(rows: list[dict[str, Any]]) -> list[str]:
        return list(dict.fromkeys(key for row in rows for key in row))

    @classmethod
    def _csv(cls, rows: list[dict[str, Any]]) -> bytes:
        output = StringIO(newline="")
        fields = cls._fieldnames(rows)
        writer = csv.DictWriter(output, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)
        return output.getvalue().encode("utf-8-sig")

    @classmethod
    def _xlsx(cls, rows: list[dict[str, Any]]) -> bytes:
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Cleaned data"
        fields = cls._fieldnames(rows)
        sheet.append(fields)
        for row in rows:
            sheet.append([row.get(field) for field in fields])
        output = BytesIO()
        workbook.save(output)
        return output.getvalue()

    @staticmethod
    def _summary_lines(detail: dict[str, Any]) -> list[tuple[str, str]]:
        summary = detail["summary"]
        return [(str(key).replace("_", " ").title(), str(value)) for key, value in summary.items()]

    @classmethod
    def _pdf(cls, detail: dict[str, Any], preview: list[dict[str, Any]]) -> bytes:
        output = BytesIO()
        document = SimpleDocTemplate(output, pagesize=A4)
        styles = getSampleStyleSheet()
        story = [Paragraph(f"Cleaning report: {detail['name']}", styles["Title"]), Spacer(1, 12)]
        story.append(Table(cls._summary_lines(detail), colWidths=(180, 300)))
        if preview:
            story.extend([Spacer(1, 16), Paragraph("Cleaned data preview", styles["Heading2"])])
            fields = cls._fieldnames(preview)[:8]
            values = [fields] + [[str(row.get(field, ""))[:40] for field in fields] for row in preview]
            table = Table(values, repeatRows=1)
            table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.25, colors.grey), ("FONTSIZE", (0, 0), (-1, -1), 7)]))
            story.append(table)
        document.build(story)
        return output.getvalue()

    @classmethod
    def _docx(cls, detail: dict[str, Any], preview: list[dict[str, Any]]) -> bytes:
        document = Document()
        document.add_heading(f"Cleaning report: {detail['name']}", 0)
        table = document.add_table(rows=0, cols=2)
        for label, value in cls._summary_lines(detail):
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value
        if preview:
            document.add_heading("Cleaned data preview", level=1)
            fields = cls._fieldnames(preview)[:8]
            preview_table = document.add_table(rows=1, cols=len(fields))
            for index, field in enumerate(fields):
                preview_table.rows[0].cells[index].text = field
            for row in preview:
                cells = preview_table.add_row().cells
                for index, field in enumerate(fields):
                    cells[index].text = str(row.get(field, ""))[:80]
        output = BytesIO()
        document.save(output)
        return output.getvalue()
